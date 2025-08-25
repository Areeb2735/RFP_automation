import streamlit as st
import os
from google import genai
from google.genai import types
from dotenv import load_dotenv
from markitdown import MarkItDown
md = MarkItDown(enable_plugins=False)
from docx import Document
from io import BytesIO
import re

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')

# Initialize GenAI client
client = genai.Client(api_key=GEMINI_API_KEY)

# Function to generate Scope of Work section
def generate_scope_of_work(project_name, scope_text, duration, criteria_text, department, example_rfp_content):
    model = "gemini-2.0-flash"
    contents = [
        types.Content(
            role="user",
            parts=[
                types.Part.from_text(text="Write a professional “SCOPE OF WORK” section for this RFP."),
            ],
        ),
    ]
    system_msg = f"""
You are a professional assistant tasked with drafting formal RFP (Request for Proposal) documentation.
Use the inputs provided below to write a well-structured and professional “SCOPE OF WORK” section for an RFP:
• Project Name: {project_name}
• Scope Description: {scope_text}
• Duration: {duration}
• Evaluation Criteria Summary: {criteria_text}
• Requesting Department: {department}

Instructions:
• Structure the section using clear headings and bullet points where appropriate.
• Write in a formal and business-like tone, suitable for government or enterprise procurement use.
• Clearly define the objectives, deliverables, services to be performed, and the expected outcomes of the project.
• Mention the duration and timeline for the execution of the work.
• If relevant, include responsibilities of the vendor and the requesting department.
• Keep the language precise, unambiguous, and action-oriented.
• DO NOT copy any content from the example RFP, but: ⚠️ Use the formatting and structural style of the provided example RFP document to guide tone, layout, and organization.

Example RFP content: {example_rfp_content}

Task: Generate a professional and comprehensive “SCOPE OF WORK” section based on the provided inputs.
"""
    config = types.GenerateContentConfig(
        response_mime_type="text/plain",
        system_instruction=[types.Part.from_text(text=system_msg)],
    )
    response = client.models.generate_content(
        model=model,
        contents=contents,
        config=config,
    )
    return response.text

def generate_creteria(project_name, scope_text, duration, criteria_text, department, example_rfp_content):
    model = "gemini-2.0-flash"
    contents = [
        types.Content(
            role="user",
            parts=[
                types.Part.from_text(text="Write a professional “CRITERIA FOR SELECTING THE WINNER” section for this RFP."),
            ],
        ),
    ]
    system_msg = f"""
You are a professional assistant tasked with drafting formal RFP (Request for Proposal) documentation.
Use the inputs provided below to write a well-structured and professional “CRITERIA FOR SELECTING THE WINNER” section for an RFP:
• Project Name: {project_name}
• Scope Description: {scope_text}
• Duration: {duration}
• Evaluation Criteria Summary: {criteria_text}
• Requesting Department: {department}

Instructions:
• Structure the section using clear headings and bullet points where appropriate.
• Write in a formal and business-like tone, suitable for government or enterprise procurement use.
• If relevant, include responsibilities of the vendor and the requesting department.
• Keep the language precise, unambiguous, and action-oriented.
• DO NOT copy any content from the example RFP, but: ⚠️ Use the formatting and structural style of the provided example RFP document to guide tone, layout, and organization.

Example RFP content: {example_rfp_content}

Task: Generate a professional and comprehensive “CRITERIA FOR SELECTING THE WINNER” section based on the provided inputs.
"""
    config = types.GenerateContentConfig(
        response_mime_type="text/plain",
        system_instruction=[types.Part.from_text(text=system_msg)],
    )
    response = client.models.generate_content(
        model=model,
        contents=contents,
        config=config,
    )
    return response.text

def create_docx(project_name, scope_text, criteria_text):
    doc = Document()
    doc.add_heading(f"RFP: {project_name}", level=1)

    doc.add_heading("Scope of Work", level=2)
    doc.add_paragraph(scope_text)

    doc.add_heading("Criteria for Selecting the Winner", level=2)
    doc.add_paragraph(criteria_text)

    doc.add_paragraph("\n*Boilerplate sections appended here*")

    # Save to in-memory file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def markdown_to_styled_docx(md_text):
    doc = Document()

    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)

        # Bold
        elif '**' in line:
            para = doc.add_paragraph()
            bold_parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in bold_parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

        # Bullets
        elif line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')

        elif line == '':
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(line)

    # Save to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.set_page_config(page_title="RFP Generator", layout="wide")
st.title("🚀 RFP Generator AI")
st.markdown("Generate tailored RFP sections in seconds using AI and your past documents.")

with st.sidebar:
    st.header("Input Parameters")
    project_name = st.text_input("Project Name", placeholder="e.g. AI Call Center Agent")
    scope_text = st.text_area("Scope of Work Description", height=150)
    duration = st.text_input("Duration", placeholder="e.g. Kickoff: July 1; Draft Review: July 15")
    department = st.text_input("Department", placeholder="e.g. IT Department")
    criteria_text = st.text_area("Criteria for Selecting the Winner", height=150, placeholder="e.g. Cost, Experience, Technical Merit")
    rfp_file = st.file_uploader("Upload Past RFP Text File", type=["pdf"], help="Use .pdf exports of your past RFPs for format guidance.")

# Read uploaded RFP content
early_rfp_content = ""
if rfp_file is not None:
    early_rfp_content = md.convert(rfp_file)
    # early_rfp_content = rfp_file.read().decode('utf-8')

# Generate button
if st.sidebar.button("Generate RFP Sections"):
    if not all([project_name, scope_text, duration, criteria_text, department]):
        st.sidebar.error("Please fill in all required fields.")
    else:
        with st.spinner("Generating Scope of Work..."):
            scope_section = generate_scope_of_work(project_name, scope_text, duration, criteria_text, department, example_rfp_content = early_rfp_content)
        st.subheader("SCOPE OF WORK")
        # st.write(scope_section)
        edited_scope = st.text_area("Edit Scope of Work Section Below", value=scope_section, height=300)

        # Placeholder for Criteria section
        with st.spinner("Generating CRITERIA FOR SELECTING THE WINNER..."):
            criteria_section = generate_creteria(project_name, scope_text, duration, criteria_text, department, example_rfp_content = early_rfp_content)
        st.subheader("CRITERIA FOR SELECTING THE WINNER")
        # st.info("Criteria generation coming soon...")
        edited_criteria = st.text_area("Edit criteria Section Below", value=criteria_section, height=300)

        # docx_file = create_docx(project_name, edited_scope, edited_criteria)

        # Export
        rfp_document = f"# RFP: {project_name}\n\n## Scope of Work\n{scope_section}\n\n## Criteria for Selecting the Winner\n...\n\n*Boilerplate sections appended here*"
        st.download_button(
            label="Download RFP as Markdown",
            data=rfp_document,
            file_name=f"RFP_{project_name.replace(' ', '_')}.md",
            mime="text/markdown"
            # mime="pdf"
        )

        docx_file = markdown_to_styled_docx(rfp_document)
        st.download_button(
            label="Download RFP as Word Document (.docx)",
            data=docx_file,
            file_name=f"RFP_{project_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


